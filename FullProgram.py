import os
import serial
import time
import mysql.connector
from datetime import datetime, timedelta
import threading
import msvcrt  # For Windows getch() function
import csv
from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
import gc
import re
# Used to print files 
import win32ui
import win32con

exit_flag = False

# Get the directory of the current script
script_dir = os.path.dirname(os.path.abspath(__file__))

# Extracts the number from a string
def extract_number(input_string):
    pattern = r"\d+(\.\d+)?"
    match = re.search(pattern, input_string)
    if match:
        extracted_number = match.group(0)
        return int(float(extracted_number))  # Convert to integer
    else:
        print("No numbers found in the input string.")
        return None

# Database connection pool
connection_pool = mysql.connector.pooling.MySQLConnectionPool(
    pool_name="mypool",
    pool_size=5,
    host="plesk.remote.ac",
    user="ws330240_Alistair",
    password="ea#4M786q",
    database="ws330240_AandR"
)

def get_db_connection():
    return connection_pool.get_connection()

def print_file_to_printer(data_string):
    printer_name = "Alistairsprinter"  # Adjust to your printer
    hdc = win32ui.CreateDC()
    hdc.CreatePrinterDC(printer_name)
    hdc.StartDoc("Weigh Head Scaled")
    hdc.StartPage()
    
    y = 100  # Starting Y position

    lines = data_string.split('\n')
    
    for line in lines:
        # Determine the font size and style based on the content of the line
        if 'BagID' in line:
            font_size = 200
            font_style = win32con.FW_BOLD
            font_italic = False
        elif 'Product' in line:
            font_size = 180
            font_style = win32con.FW_NORMAL
            font_italic = False
        else:
            font_size = 160
            font_style = win32con.FW_NORMAL
            font_italic = False

        # Create and select the font
        font = win32ui.CreateFont({
            "name": "Arial",  # Use the desired font name
            "height": font_size,
            "weight": font_style,
            "italic": font_italic,
        })
        hdc.SelectObject(font)
        
        # Write the actual line of text
        hdc.TextOut(100, y, line.strip())
        y += font_size * 10  # Increment Y position for next line based on font size

    hdc.EndPage()
    hdc.EndDoc()
    hdc.DeleteDC()
    pass

def print_file_to_word_doc(file_path, report_date):
    doc = Document()
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'  # Apply table style

    # Add header row with column names
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'BagID'
    hdr_cells[1].text = 'GrossWeight'
    hdr_cells[2].text = 'BatchNumb'
    hdr_cells[3].text = 'ProductType'
    hdr_cells[4].text = 'DateandTime'

    try:
        with open(file_path, 'r') as file:
            csv_reader = csv.reader(file)
            next(csv_reader)  # Skip header row
            for row in csv_reader:
                row_cells = table.add_row().cells
                for i, cell_value in enumerate(row):
                    row_cells[i].text = cell_value
    except FileNotFoundError:
        print(f"File '{file_path}' not found.")
        return
    except Exception as e:
        print(f"An error occurred while reading the file: {e}")
        return

    # Format table
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)  # Adjust font size

    # Center align text in all cells
    for row in table.rows:
        for cell in row.cells:
            cell.paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER

    # Save Word document with date and time in the script directory
    doc_filename = f'report_{report_date.strftime("%Y-%m-%d_%H-%M-%S")}.docx'
    doc.save(os.path.join(script_dir, doc_filename))  
    print(f"Word document saved to {os.path.join(script_dir, doc_filename)}")
    pass

def read_serial_data(port, baud_rate, ser=None):
    global exit_flag
    if ser is None:
        ser = serial.Serial(port, baud_rate)  # Open the serial port
        print(f"Opened serial port {port}")

    try:
        read_count = 0  
        batch_number = get_batch_number()
        if batch_number is None:  
            return

        product = get_product_type()
        if product is None: 
            return

        while True:
            if exit_flag:
                print("Exiting to menu...")
                break  # Exit loop

            if ser.in_waiting > 0:
                data = ser.read(ser.in_waiting).decode("utf-8")
                lines = data.split('\n')
                
                for decoded_data in lines:
                    if decoded_data.startswith("Gross"):
                        number = extract_number(decoded_data)
                        read_count += 1

                        if read_count % 22 == 0:
                            batch_number = get_batch_number()
                            product = get_product_type()

                        current_time = datetime.now()
                        mydb = get_db_connection()
                        mycursor = mydb.cursor()
                        sql = "INSERT INTO `LakesWeighHead` (`BagID`, `GrossWeight`, `DateandTime`, `BatchNumb`, `ProductType`) VALUES ('', %s, %s, %s, %s)"
                        val = (number, current_time, batch_number, product)
                        mycursor.execute(sql, val)
                        mydb.commit()

                        idsql = "SELECT `BagID` FROM `LakesWeighHead` ORDER BY `BagID` DESC LIMIT 1;"
                        mycursor.execute(idsql)
                        CurrentID = mycursor.fetchone()
                        most_recent_id = CurrentID[0]

                        print(f"\n\nBatch: {batch_number}   Weight: {number}    Product: {product}    BagID: {most_recent_id}       Date and time: {current_time}\n")

                        data_string = f"\n\nBatch: {batch_number}   Weight: {number}    Product: {product}    BagID: {most_recent_id}"
                        print_file_to_printer(data_string)
                
                        mycursor.close()
                        mydb.close()

            time.sleep(0.01)  # Small delay to prevent CPU overload

    except serial.SerialException as e:
        print(f"Serial exception: {e}")

    except KeyboardInterrupt:
        print("Interrupted by user. Exiting to menu.")

    finally:
        if ser and ser.is_open:
            ser.flush()  # Flush the serial port
            ser.close()  # Close the serial connection
            del ser
            gc.collect()
            time.sleep(2)  # Wait for 2 seconds before re-opening
            print(f"Closed serial port {port}")

            # Reopen menu
            menu(port, baud_rate)

def get_batch_number():
    while True:
        user_input = input("Please enter the batch number to begin: ")
        if user_input == '0':
            return None
        try:
            return int(user_input)
        except ValueError:
            print("Invalid input. Please enter an integer.")

def get_product_type():
    while True:
        try:
            product_input = int(input("""Enter the Product Type:\n
            1) Skins\n
            2) Sticks\n"""))
            if product_input == 1:
                return "Skins"
            elif product_input == 2:
                return "Sticks"
            else:
                print("Invalid input. Please enter 1 or 2.")
        except ValueError:
            print("Invalid input. Please enter an integer.")

def update(port, baud_rate, ser):
    print("Update is used if a bag is broken or incorrectly weighed\n")
    print("Please Enter the Bag ID for the bag that needs reweighing\n")
    
    while True:
        bag_id = input("-->")
        if bag_id.isdigit():
            break
        else:
            print("Invalid input. Please enter a numeric Bag ID.")

    try:
        while True:
            if ser.in_waiting > 0:
                data = ser.readline().strip()
                decoded_data = data.decode("utf-8")

                if decoded_data.startswith("Gross"):
                    new_weight = extract_number(decoded_data)

                    if new_weight is not None:
                        mydb = get_db_connection()
                        mycursor = mydb.cursor()
                        sql = "UPDATE `LakesWeighHead` SET `GrossWeight` = %s WHERE `BagID` = %s"
                        val = (new_weight, bag_id)
                        mycursor.execute(sql, val)
                        mydb.commit()

                        sql = "SELECT `BatchNumb`, `ProductType` FROM `LakesWeighHead` WHERE `BagID` = %s"
                        val = (bag_id,)
                        mycursor.execute(sql, val)
                        result = mycursor.fetchone()

                        if result:
                            batch_number, product = result
                            current_time = datetime.now()
                            print(f"Updated Bag ID: {bag_id}, New Weight: {new_weight}, Batch Number: {batch_number}, Product: {product}, Date and Time: {current_time}")

                            data_string = f"\n\nBatch: {batch_number}   Weight: {new_weight}    Product: {product}    BagID: {bag_id}"
                            print_file_to_printer(data_string)

                            break  # Exit after a successful update
                        else:
                            print("Bag ID not found. Please try again.")

                        mycursor.close()
                        mydb.close()
                    else:
                        print("Failed to extract a valid number from the scale data. Please try again.")

            time.sleep(0.1)

    except serial.SerialException as e:
        print(f"Serial exception: {e}")

    finally:
        menu(port, baud_rate, ser)  # Return to the main menu
    pass

def previous_day_reports(port, baud_rate, ser):
    today = datetime.now().date()
    days = [today - timedelta(days=i) for i in range(20)]  # Last 20 days including today

    print("Select a day to view the report:")
    for i, day in enumerate(days, 1):
        print(f"{i}) {day}")

    while True:
        user_input = input("Enter the number corresponding to the day: ")
        if user_input == '0':
            menu(port, baud_rate)
            return
        try:
            selected_day = int(user_input)
            if 1 <= selected_day <= 20:
                selected_date = days[selected_day - 1]
                break
            else:
                print("Invalid input. Please enter a number between 1 and 20.")
        except ValueError:
            print("Invalid input. Please enter a number.")

    next_day = selected_date + timedelta(days=1)
    
    mydb = get_db_connection()
    mycursor = mydb.cursor()

    sql = "SELECT * FROM `LakesWeighHead` WHERE `DateandTime` >= %s AND `DateandTime` < %s"
    val = (selected_date, next_day)
    mycursor.execute(sql, val)

    results = mycursor.fetchall()
    
    if results:
        csv_filename = f"previous_day_report_{selected_date}.csv"  # CSV filename based on selected date
        csv_file_path = os.path.join(script_dir, csv_filename)
        with open(csv_file_path, mode='w', newline='') as csvfile:
            csv_writer = csv.writer(csvfile)
            csv_writer.writerow(["BagID", "GrossWeight", "BatchNumb", "ProductType", "DateandTime"])  # Write header
            for row in results:
                csv_writer.writerow(row)  
        print(f"Previous day report saved to {csv_file_path}.")
        
        print_file_to_word_doc(csv_file_path, selected_date)
        
    else:
        print(f"No entries found for {selected_date}.")

    mycursor.close()
    mydb.close()

    menu(port, baud_rate, ser)
    pass


def exit_listener(port, baud_rate):
    global exit_flag
    print("Press '0' to exit to the menu \n")
    while True:
        if msvcrt.kbhit():  
            input_char = msvcrt.getch().decode('utf-8')
            if input_char == '0':
                exit_flag = True
                return

def menu(port, baud_rate, ser=None):
    print("\nWeigh Head System\n")
    print("1) Start Printing\n")
    print("2) Update a Field\n")
    print("3) Reports\n")

    if ser is None:
        ser = serial.Serial(port, baud_rate)  # Open the serial port
        print(f"Opened serial port {port}")

    MenuOption = input("-->")
    if MenuOption == "1":
        global exit_flag
        exit_flag = False
        exit_thread = threading.Thread(target=exit_listener, args=(port, baud_rate))
        exit_thread.daemon = True
        exit_thread.start()
        read_serial_data(port, baud_rate, ser)
    elif MenuOption == "2":
        update(port, baud_rate, ser)
    elif MenuOption == "3":
        previous_day_reports(port, baud_rate, ser)
    else:
        print("Invalid Input")
        menu(port, baud_rate, ser)  # Return to menu for invalid input

if __name__ == "__main__":
    port = 'COM3'  # Set your COM port
    baud_rate = 9600  # Set the baud rate

    menu(port, baud_rate)

    input("Press Enter to exit...")
