import os
import mysql.connector
from datetime import datetime, timedelta
import csv
from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT

# Get the directory of the current script
script_dir = os.path.dirname(os.path.abspath(__file__))

# Database connection pool
connection_pool = mysql.connector.pooling.MySQLConnectionPool(
    pool_name="mypool",
    pool_size=5,
    host="plesk.remote.ac",
    user="ws330240_Alistair",
    password="ea#4M786q",
    database="ws330240_AandR"
)

def get_db_connection():
    return connection_pool.get_connection()

def print_file_to_word_doc(file_path, report_date):
    doc = Document()
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'  # Apply table style

    # Add header row with column names
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'BagID'
    hdr_cells[1].text = 'GrossWeight'
    hdr_cells[2].text = 'BatchNumb'
    hdr_cells[3].text = 'ProductType'
    hdr_cells[4].text = 'DateandTime'

    try:
        with open(file_path, 'r') as file:
            csv_reader = csv.reader(file)
            next(csv_reader)  # Skip header row
            for row in csv_reader:
                row_cells = table.add_row().cells
                for i, cell_value in enumerate(row):
                    row_cells[i].text = cell_value
    except FileNotFoundError:
        print(f"File '{file_path}' not found.")
        return
    except Exception as e:
        print(f"An error occurred while reading the file: {e}")
        return

    # Format table
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)  # Adjust font size

    # Center align text in all cells
    for row in table.rows:
        for cell in row.cells:
            cell.paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER

    # Save Word document with date and time in the script directory
    doc_filename = f'report_{report_date.strftime("%Y-%m-%d_%H-%M-%S")}.docx'
    doc.save(os.path.join(script_dir, doc_filename))  
    print(f"Word document saved to {os.path.join(script_dir, doc_filename)}")

def previous_day_reports():
    today = datetime.now().date()
    days = [today - timedelta(days=i) for i in range(20)]  # Last 20 days including today

    print("Select a day to view the report:")
    for i, day in enumerate(days, 1):
        print(f"{i}) {day}")

    while True:
        user_input = input("Enter the number corresponding to the day: ")
        if user_input == '0':
            return
        try:
            selected_day = int(user_input)
            if 1 <= selected_day <= 20:
                selected_date = days[selected_day - 1]
                break
            else:
                print("Invalid input. Please enter a number between 1 and 20.")
        except ValueError:
            print("Invalid input. Please enter a number.")

    next_day = selected_date + timedelta(days=1)
    
    mydb = get_db_connection()
    mycursor = mydb.cursor()

    sql = "SELECT * FROM `LakesWeighHead` WHERE `DateandTime` >= %s AND `DateandTime` < %s"
    val = (selected_date, next_day)
    mycursor.execute(sql, val)

    results = mycursor.fetchall()
    
    if results:
        csv_filename = f"previous_day_report_{selected_date}.csv"  # CSV filename based on selected date
        csv_file_path = os.path.join(script_dir, csv_filename)
        with open(csv_file_path, mode='w', newline='') as csvfile:
            csv_writer = csv.writer(csvfile)
            csv_writer.writerow(["BagID", "GrossWeight", "BatchNumb", "ProductType", "DateandTime"])  # Write header
            for row in results:
                csv_writer.writerow(row)  
        print(f"Previous day report saved to {csv_file_path}.")
        
        print_file_to_word_doc(csv_file_path, selected_date)
        
    else:
        print(f"No entries found for {selected_date}.")

    mycursor.close()
    mydb.close()

if __name__ == "__main__":
    previous_day_reports()
    input("Press Enter to exit...")
