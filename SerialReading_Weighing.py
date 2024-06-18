import os
import serial
import time
import mysql.connector
from datetime import datetime, timedelta
import threading
import msvcrt  # For Windows getch() function
import csv
from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
import gc
import re
# Used to print files 
import win32ui
import win32con

exit_flag = False

# Get the directory of the current script
script_dir = os.path.dirname(os.path.abspath(__file__))

# Extracts the number from a string
def extract_number(input_string):
    pattern = r"\d+(\.\d+)?"
    match = re.search(pattern, input_string)
    if match:
        extracted_number = match.group(0)
        return int(float(extracted_number))  # Convert to integer
    else:
        print("No numbers found in the input string.")
        return None

# Database connection pool
connection_pool = mysql.connector.pooling.MySQLConnectionPool(
    pool_name="mypool",
    pool_size=5,
    host="",
    user="",
    password="",
    database=""
)

def get_db_connection():
    return connection_pool.get_connection()

def print_file_to_printer(data_string):
    try:
        printer_name = ""  # Adjust to your printer
        hdc = win32ui.CreateDC()
        hdc.CreatePrinterDC(printer_name)
        hdc.StartDoc("Weigh Head Scaled")
        hdc.StartPage()
        
        y = 100  # Starting Y position

        lines = data_string.split('\n')
        
        for line in lines:
            # Split the line further if necessary
            sublines = line.split('    ')  # Split by multiple spaces

            for subline in sublines:
                # Determine the font size and style based on the content of the line
                if 'BagID' in subline:
                    font_size = 500
                    font_style = win32con.FW_BOLD
                    font_italic = False
                elif 'Product' in subline:
                    font_size = 480
                    font_style = win32con.FW_NORMAL
                    font_italic = False
                else:
                    font_size = 460
                    font_style = win32con.FW_NORMAL
                    font_italic = False

                # Create and select the font
                font = win32ui.CreateFont({
                    "name": "Arial",  # Use the desired font name
                    "height": font_size,
                    "weight": font_style,
                    "italic": font_italic,
                })
                hdc.SelectObject(font)
                
                # Write the actual line of text
                hdc.TextOut(100, y, subline.strip())
                y += int(font_size * 1.2)  # Increment Y position for next line based on font size

        hdc.EndPage()
        hdc.EndDoc()
        hdc.DeleteDC()
    except Exception as e:
        print(f"An error occurred: Unable to open printer ({e})")


        
def print_file_to_word_doc(file_path, report_date):
    doc = Document()
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'  # Apply table style

    # Add header row with column names
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'BagID'
    hdr_cells[1].text = 'GrossWeight'
    hdr_cells[2].text = 'BatchNumb'
    hdr_cells[3].text = 'ProductType'
    hdr_cells[4].text = 'DateandTime'

    try:
        with open(file_path, 'r') as file:
            csv_reader = csv.reader(file)
            next(csv_reader)  # Skip header row
            for row in csv_reader:
                row_cells = table.add_row().cells
                for i, cell_value in enumerate(row):
                    row_cells[i].text = cell_value
    except FileNotFoundError:
        print(f"File '{file_path}' not found.")
        return
    except Exception as e:
        print(f"An error occurred while reading the file: {e}")
        return

    # Format table
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)  # Adjust font size

    # Center align text in all cells
    for row in table.rows:
        for cell in row.cells:
            cell.paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER

    # Save Word document with date and time in the script directory
    doc_filename = f'report_{report_date.strftime("%Y-%m-%d_%H-%M-%S")}.docx'
    doc.save(os.path.join(script_dir, doc_filename))  
    print(f"Word document saved to {os.path.join(script_dir, doc_filename)}")
    pass

def read_serial_data(port, baud_rate, ser=None):
    print("During the Weighing process press 0 at any point to return to the main menu")
    global exit_flag
    if ser is None:
        ser = serial.Serial(port, baud_rate)  # Open the serial port
        print(f"Opened serial port {port}")

    try:
        read_count = 0  
        batch_number = get_batch_number()
        if batch_number is None:  
            return

        product = get_product_type()
        if product is None: 
            return

        while True:
            if exit_flag:
                print("Exiting to menu...")
                break  # Exit loop

            if ser.in_waiting > 0:
                data = ser.read(ser.in_waiting).decode("utf-8")
                lines = data.split('\n')
                
                for decoded_data in lines:
                    if decoded_data.startswith("Gross"):
                        number = extract_number(decoded_data)
                        read_count += 1

                        if read_count % 22 == 0:
                            batch_number = get_batch_number()
                            product = get_product_type()

                        current_time = datetime.now()
                        mydb = get_db_connection()
                        mycursor = mydb.cursor()
                        sql = "INSERT INTO `LakesWeighHead` (`BagID`, `GrossWeight`, `DateandTime`, `BatchNumb`, `ProductType`) VALUES ('', %s, %s, %s, %s)"
                        val = (number, current_time, batch_number, product)
                        mycursor.execute(sql, val)
                        mydb.commit()

                        idsql = "SELECT `BagID` FROM `LakesWeighHead` ORDER BY `BagID` DESC LIMIT 1;"
                        mycursor.execute(idsql)
                        CurrentID = mycursor.fetchone()
                        most_recent_id = CurrentID[0]

                        print(f"\n\nBatch: {batch_number}   Weight: {number}    Product: {product}    BagID: {most_recent_id}       Date and time: {current_time}\n")

                        data_string = f"\n\nBatch: {batch_number}   Weight: {number}    Product: {product}    BagID: {most_recent_id}"
                        print_file_to_printer(data_string)
                
                        mycursor.close()
                        mydb.close()

            time.sleep(0.01)  # Small delay to prevent CPU overload

    except serial.SerialException as e:
        print(f"Serial exception: {e}")

    except KeyboardInterrupt:
        print("Interrupted by user. Exiting to menu.")

    except Exception as e:
        print(f"An error occurred: {e}")

    finally:
        if ser and ser.is_open:
            ser.flush()  # Flush the serial port
            ser.close()  # Close the serial connection
            del ser
            gc.collect()
            time.sleep(2)  # Wait for 2 seconds before re-opening
            print(f"Closed serial port {port}")

            # Reopen menu
            menu(port, baud_rate)

def get_batch_number():
    while True:
        user_input = input("Please enter the batch number to begin: ")
        if user_input == '0':
            return None
        try:
            return int(user_input)
        except ValueError:
            print("Invalid input. Please enter an integer.")

def get_product_type():
    while True:
        try:
            product_input = int(input("""Enter the Product Type:\n
            1) Skins\n
            2) Sticks\n"""))
            if product_input == 1:
                return "Skins"
            elif product_input == 2:
                return "Sticks"
            else:
                print("Invalid input. Please enter 1 or 2.")
        except ValueError:
            print("Invalid input. Please enter an integer.")

def update(port, baud_rate, ser):
    print("Update is used if a bag is broken or incorrectly weighed\n")
    print("Please Enter the Bag ID for the bag that needs reweighing\n")
    
    while True:
        bag_id = input("-->")
        if bag_id.isdigit():
            break
        else:
            print("Invalid input. Please enter a numeric Bag ID.")

    try:
        ser = serial.Serial(port, baud_rate) if ser is None else ser
        print(f"Listening for reweigh data on {port}...")

        while True:
            if ser.in_waiting > 0:
                data = ser.read(ser.in_waiting).decode("utf-8")
                lines = data.split('\n')
                
                for decoded_data in lines:
                    if decoded_data.startswith("Gross"):
                        new_weight = extract_number(decoded_data)
                        current_time = datetime.now()

                        mydb = get_db_connection()
                        mycursor = mydb.cursor()

                        # Retrieve the current batch number and product type
                        select_sql = "SELECT BatchNumb, ProductType FROM `LakesWeighHead` WHERE `BagID` = %s"
                        mycursor.execute(select_sql, (bag_id,))
                        result = mycursor.fetchone()
                        batch_number = result[0]
                        product_type = result[1]

                        update_sql = "UPDATE `LakesWeighHead` SET `GrossWeight` = %s, `DateandTime` = %s WHERE `BagID` = %s"
                        val = (new_weight, current_time, bag_id)
                        mycursor.execute(update_sql, val)
                        mydb.commit()
                        mycursor.close()
                        mydb.close()

                        print(f"Updated BagID {bag_id} with new weight: {new_weight}\n")
                        
                        data_string = f"BagID: {bag_id}\nWeight: {new_weight}\nProduct: {product_type}\nBatch: {batch_number} *"
                        print_file_to_printer(data_string)

                        return

            if exit_flag:
                print("Exiting update mode...")
                break  # Exit loop

            time.sleep(0.01)  # Small delay to prevent CPU overload

    except serial.SerialException as e:
        print(f"Serial exception: {e}")

    except KeyboardInterrupt:
        print("Interrupted by user. Exiting to menu.")

    except Exception as e:
        print(f"An error occurred: {e}")

    finally:
        if ser and ser.is_open:
            ser.flush()  # Flush the serial port
            ser.close()  # Close the serial connection
            del ser
            gc.collect()
            time.sleep(2)  # Wait for 2 seconds before re-opening
            print(f"Closed serial port {port}")

            # Reopen menu
            menu(port, baud_rate)



def monitor_keyboard_input():
    global exit_flag
    while True:
        if msvcrt.kbhit():
            key = msvcrt.getch()
            if key == b'0':
                exit_flag = True
                break

def menu(port, baud_rate):
    global exit_flag
    ser = None

    while True:
        try:
            choice = input("""Please make your selection:\n
            1) Start Weighing\n
            2) Update a Weight\n
            5) End program\n""")
            
            if choice == '1':
                exit_flag = False
                ser = serial.Serial(port, baud_rate) if ser is None else ser
                threading.Thread(target=monitor_keyboard_input, daemon=True).start()
                read_serial_data(port, baud_rate, ser)
            elif choice == '2':
                exit_flag = False
                ser = serial.Serial(port, baud_rate) if ser is None else ser
                update(port, baud_rate, ser)
            elif choice == '[][]':
                file_path = input("Please enter the file path of the CSV file: ")
                report_date_str = input("Please enter the report date (YYYY-MM-DD): ")
                try:
                    report_date = datetime.strptime(report_date_str, '%Y-%m-%d')
                    print_file_to_word_doc(file_path, report_date)
                except ValueError:
                    print("Invalid date format. Please enter the date in YYYY-MM-DD format.")
            elif choice == '[][':
                exit_flag = True
                if ser and ser.is_open:
                    ser.flush()  # Flush the serial port
                    ser.close()  # Close the serial connection
                    del ser
                    gc.collect()
                    time.sleep(2)  # Wait for 2 seconds before re-opening
                    print(f"Closed serial port {port}")
                # Call the main menu function here if available
            elif choice == '5':
                exit_flag = True
                if ser and ser.is_open:
                    ser.flush()  # Flush the serial port
                    ser.close()  # Close the serial connection
                    del ser
                    gc.collect()
                    time.sleep(2)  # Wait for 2 seconds before re-opening
                    print(f"Closed serial port {port}")
                print("Exiting program...")
                exit(0)
            else:
                print("Invalid input. Please enter a number from 1 to 5.")
        except KeyboardInterrupt:
            exit_flag = True
            if ser and ser.is_open:
                ser.flush()  # Flush the serial port
                ser.close()  # Close the serial connection
                del ser
                gc.collect()
                time.sleep(2)  # Wait for 2 seconds before re-opening
                print(f"Closed serial port {port}")
            print("Interrupted by user. Exiting program...")
            exit(0)

if __name__ == "__main__":
    serial_port = "COM2"  # Replace with your serial port
    baud_rate = 9600  # Replace with your baud rate
    menu(serial_port, baud_rate)
